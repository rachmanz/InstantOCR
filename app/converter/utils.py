import json
import csv
import io
from docx import Document
import pandas as pd

class FormatConverter:
    @staticmethod
    def to_txt(text: str) -> bytes:
        return text.encode('utf-8')

    @staticmethod
    def to_json(text: str) -> bytes:
        data = {"status": "success", "extracted_text": text}
        return json.dumps(data, indent=2).encode('utf-8')

    @staticmethod
    def to_docx(text: str) -> bytes:
        doc = Document()
        doc.add_heading('Instant-OCR Result', level=1)
        doc.add_paragraph(text)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def to_csv(text: str) -> bytes:
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        writer.writerow(["Line", "Content"])
        
        for idx, line in enumerate(text.splitlines(), start=1):
            if line.strip():
                writer.writerow([idx, line.strip()])
                
        return buffer.getvalue().encode('utf-8')

    @staticmethod
    def to_xlsx(text: str) -> bytes:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        df = pd.DataFrame({"Extracted Text": lines})
        
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name="OCR Output")
            
        return buffer.getvalue()