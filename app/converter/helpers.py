import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def generate_docx(text):
    """Membuat file Word (.docx) di dalam memori (BytesIO)"""
    doc = Document()
    doc.add_heading('Hasil OCR Vision', level=1)
    
    # Masukkan teks per paragraf
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf(text):
    """Membuat file PDF di dalam memori (BytesIO)"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom style agar teks rapi
    body_style = ParagraphStyle(
        'OCRBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=8
    )
    
    story = [
        Paragraph("<b>Hasil Ekstraksi OCR</b>", styles['Title']),
        Spacer(1, 12)
    ]
    
    # Masukkan baris teks
    for line in text.split('\n'):
        if line.strip():
            story.append(Paragraph(line, body_style))
            
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_txt(text):
    """Membuat file teks (.txt) di dalam memori (BytesIO)"""
    buffer = io.BytesIO()
    buffer.write(text.encode('utf-8'))
    buffer.seek(0)
    return buffer